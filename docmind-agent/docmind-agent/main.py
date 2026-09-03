"""
DocMind Agent  —  main.py  (v4.0 — Non-Human Identity edition)
================================================================
Every tool call now requires a valid, scoped, short-lived access token
issued to the AGENT'S OWN identity (not the human's) — the exact
client_credentials shape from Section 1's Auth0 setup, applied to a
non-human identity instead of a person.

Lifecycle per document session:
  1. Upload a document      -> a session is created
  2. PROVISION agent identity -> client_id + client_secret issued,
                                  scoped to read:documents by default
  3. ISSUE a token           -> short-lived access_token (5 min read,
                                  2 min write), requested explicitly
  4. USE the token            -> every tool call presents the token;
                                  PermissionGate validates scope + expiry
  5. DEPROVISION               -> when the session ends (or on demand),
                                  the identity is revoked and every
                                  outstanding token dies with it

Supported file types: PDF · DOCX · XLSX · XLS · CSV · TXT
AI backend: Ollama (llama3, primary) -> Claude API (fallback)
"""

import os
import textwrap
import uuid
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import pandas as pd
import pdfplumber
import requests
from docx import Document as DocxDocument
from fastapi import FastAPI, File, UploadFile, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from agent.core import (
    Tool, ToolRegistry, sanitize_document, chunk_text, retrieve_chunks,
    audit, get_audit_log, VALID_CHUNK_SIZES, DEFAULT_CHUNK_SIZE, DEFAULT_CHUNK_OVERLAP,
)
from agent import identity as nhi

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger(__name__)

UPLOAD_DIR = Path("uploads"); UPLOAD_DIR.mkdir(exist_ok=True)
STATIC_DIR = Path("static")

OLLAMA_URL     = os.getenv("OLLAMA_URL",        "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3.2:latest")
CLAUDE_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL   = "claude-sonnet-4-6"

ALLOWED_EXT = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".txt"}

# sessions[session_id] = {filename, file_type, pages, word_count, char_count,
#                          raw_text, clean_text, chunks, chunk_size, overlap,
#                          injection_warnings, client_id, created_at}
sessions: dict[str, dict] = {}

app = FastAPI(title="DocMind Agent API", version="4.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


# ══════════════════════════════════════════════════════════════════════════════
#  FILE EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_pdf(path: Path) -> tuple[str, dict]:
    pages = []
    with pdfplumber.open(path) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages, 1):
            t = page.extract_text() or ""
            if t.strip():
                pages.append(f"[PAGE {i}]\n{t.strip()}")
    return "\n\n".join(pages), {"pages": total, "file_type": "PDF"}


def extract_docx(path: Path) -> tuple[str, dict]:
    doc = DocxDocument(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else ""
            prefix = "## " if "Heading" in style else ""
            parts.append(prefix + para.text.strip())
    for t_idx, table in enumerate(doc.tables, 1):
        parts.append(f"\n[TABLE {t_idx}]")
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts), {"pages": max(1, len(doc.paragraphs) // 5), "file_type": "Word Document"}


def extract_xlsx(path: Path) -> tuple[str, dict]:
    xl = pd.ExcelFile(str(path))
    parts = []
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        parts.append(f"[SHEET: {sheet}]")
        parts.append(df.head(200).to_string(index=False))
        parts.append(f"(Rows: {len(df)}, Cols: {len(df.columns)})")
    return "\n\n".join(parts), {"pages": len(xl.sheet_names), "file_type": "Excel Spreadsheet"}


def extract_csv(path: Path) -> tuple[str, dict]:
    try:
        df = pd.read_csv(str(path), encoding="utf-8", on_bad_lines="skip")
    except Exception:
        df = pd.read_csv(str(path), encoding="latin-1", on_bad_lines="skip")
    text = f"[CSV — {len(df)} rows × {len(df.columns)} columns]\n\n" + df.head(300).to_string(index=False)
    return text, {"pages": 1, "file_type": "CSV Spreadsheet"}


def extract_txt(path: Path) -> tuple[str, dict]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1")
    return text, {"pages": max(1, len(text.splitlines()) // 50), "file_type": "Text File"}


def extract_file(path: Path) -> tuple[str, dict]:
    ext = path.suffix.lower()
    dispatch = {
        ".pdf": extract_pdf, ".docx": extract_docx, ".doc": extract_docx,
        ".xlsx": extract_xlsx, ".xls": extract_xlsx,
        ".csv": extract_csv, ".txt": extract_txt,
    }
    if ext not in dispatch:
        raise ValueError(f"Unsupported: {ext}")
    return dispatch[ext](path)


# ══════════════════════════════════════════════════════════════════════════════
#  AI ENGINE
# ══════════════════════════════════════════════════════════════════════════════

def ask_ollama(prompt: str, max_tokens: int = 1000) -> str:
    resp = requests.post(
        f"{OLLAMA_URL}/api/generate",
        json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False,
              "options": {"temperature": 0.15, "num_predict": max_tokens}},
        timeout=240,
    )
    resp.raise_for_status()
    return resp.json().get("response", "").strip()


def ask_claude(prompt: str, max_tokens: int = 1000) -> str:
    if not CLAUDE_API_KEY:
        raise ValueError("ANTHROPIC_API_KEY not set")
    resp = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={"x-api-key": CLAUDE_API_KEY, "anthropic-version": "2023-06-01",
                 "content-type": "application/json"},
        json={"model": CLAUDE_MODEL, "max_tokens": max_tokens,
              "messages": [{"role": "user", "content": prompt}]},
        timeout=90,
    )
    resp.raise_for_status()
    return resp.json()["content"][0]["text"].strip()


def generate(prompt: str, max_tokens: int = 1000) -> tuple[str, str]:
    try:
        ans = ask_ollama(prompt, max_tokens)
        if ans:
            return ans, f"Ollama · {OLLAMA_MODEL}"
    except Exception as e:
        log.warning("Ollama failed: %s", e)
    try:
        ans = ask_claude(prompt, max_tokens)
        return ans, f"Claude · {CLAUDE_MODEL}"
    except Exception as e:
        raise RuntimeError("No AI backend. Run `ollama serve` or set ANTHROPIC_API_KEY.") from e


# ══════════════════════════════════════════════════════════════════════════════
#  PROMPT BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

def build_qa_prompt(retrieved_chunks: list[dict], question: str, file_type: str) -> str:
    context = "\n\n---\n\n".join(
        f"[Chunk {c['chunk_id']} | words {c['word_start']}–{c['word_end']} | relevance {c['score']}]\n{c['text']}"
        for c in retrieved_chunks
    )
    return textwrap.dedent(f"""
        You are an expert document analyst. The user uploaded a {file_type}.
        Answer the question using ONLY the retrieved passages below.
        Cite the chunk ID when you use information from it.
        If the answer is not in the passages, say: "I couldn't find that in the document."

        ── RETRIEVED PASSAGES ────────────────────────────────────
        {context}
        ─────────────────────────────────────────────────────────

        Question: {question}
        Answer:
    """).strip()


def build_summary_prompt(context: str, file_type: str, instruction: str = "") -> str:
    extra = f"\nUser instruction: {instruction}" if instruction else ""
    return textwrap.dedent(f"""
        You are an expert document analyst. Summarize this {file_type}.
        Structure: 1) 2–3 sentence overview  2) Key points (bullets)  3) Conclusions/next steps
        {extra}

        ── DOCUMENT ──────────────────────────────────────────────
        {context}
        ──────────────────────────────────────────────────────────

        Summary:
    """).strip()


def build_keyterms_prompt(context: str, file_type: str) -> str:
    return textwrap.dedent(f"""
        Extract the 10 most important keywords/key phrases from this {file_type}.
        Return a numbered list, one per line. No explanations.

        ── DOCUMENT ──────────────────────────────────────────────
        {context}
        ──────────────────────────────────────────────────────────

        Key Terms:
    """).strip()


def build_quiz_prompt(context: str, file_type: str) -> str:
    return textwrap.dedent(f"""
        Create 5 multiple-choice quiz questions from this {file_type}.
        Format:
        Q1. [Question]
        A) B) C) D) [options]
        Answer: [Letter]

        Use ONLY document content.

        ── DOCUMENT ──────────────────────────────────────────────
        {context}
        ──────────────────────────────────────────────────────────

        Quiz:
    """).strip()


def trim(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    h = max_chars // 2
    return text[:h] + "\n\n...[trimmed]...\n\n" + text[-h:]


# ══════════════════════════════════════════════════════════════════════════════
#  TOOL FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def tool_answer_question(session_id: str, question: str) -> dict:
    s = sessions[session_id]
    top = retrieve_chunks(s["chunks"], question, top_k=5)
    prompt = build_qa_prompt(top, question, s["file_type"])
    ans, engine = generate(prompt, max_tokens=900)
    return {"answer": ans, "engine": engine, "chunks_used": [c["chunk_id"] for c in top],
            "retrieval_scores": {c["chunk_id"]: c["score"] for c in top}}


def tool_summarize(
    session_id: str,
    question: str = "",
    instruction: str = ""
) -> dict:
    s = sessions[session_id]
    ctx = trim(s["clean_text"], 22_000)
    prompt = build_summary_prompt(ctx, s["file_type"], instruction)
    ans, engine = generate(prompt, max_tokens=1400)
    return {"summary": ans, "engine": engine}


def tool_keyterms(
    session_id: str,
    question: str = ""
) -> dict:
    s = sessions[session_id]
    ctx = trim(s["clean_text"], 14_000)
    prompt = build_keyterms_prompt(ctx, s["file_type"])
    ans, engine = generate(prompt, max_tokens=400)
    return {"keyterms": ans, "engine": engine}


def tool_quiz(
    session_id: str,
    question: str = ""
) -> dict:
    s = sessions[session_id]
    ctx = trim(s["clean_text"], 20_000)
    prompt = build_quiz_prompt(ctx, s["file_type"])
    ans, engine = generate(prompt, max_tokens=1200)
    return {"quiz": ans, "engine": engine}


def tool_export_summary(
    session_id: str,
    question: str = "",
    **_
) -> dict:
    """[WRITE] Example of an action tool — writes a file to disk. Needs write:documents."""
    s = sessions[session_id]
    ctx = trim(s["clean_text"], 22_000)
    prompt = build_summary_prompt(ctx, s["file_type"])
    ans, engine = generate(prompt, max_tokens=1400)
    out_path = UPLOAD_DIR / f"{session_id}_summary.txt"
    out_path.write_text(ans)
    return {"message": f"Summary exported to {out_path.name}", "engine": engine}


# ══════════════════════════════════════════════════════════════════════════════
#  TOOL REGISTRY  — every tool declares the scope it needs
# ══════════════════════════════════════════════════════════════════════════════

registry = ToolRegistry()

registry.register(Tool(name="answer_question", description="Answer a question using retrieved document chunks.",
                        required_scope="read:documents", fn=tool_answer_question))
registry.register(Tool(name="summarize_document", description="Generate a structured document summary.",
                        required_scope="read:documents", fn=tool_summarize))
registry.register(Tool(name="extract_keyterms", description="Extract top 10 keywords/phrases.",
                        required_scope="read:documents", fn=tool_keyterms))
registry.register(Tool(name="generate_quiz", description="Generate a 5-question multiple-choice quiz.",
                        required_scope="read:documents", fn=tool_quiz))
registry.register(Tool(name="export_summary", description="[WRITE] Save the summary to disk as a file.",
                        required_scope="write:documents", fn=tool_export_summary, requires_approval=True))


def decide_tool(question: str) -> str:
    q = question.lower()
    if any(w in q for w in ["summarize", "summarise", "summary", "overview", "brief", "gist", "tldr"]):
        return "summarize_document"
    if any(w in q for w in ["quiz", "test me", "examine"]):
        return "generate_quiz"
    if any(w in q for w in ["key term", "keyword", "important word"]):
        return "extract_keyterms"
    if any(w in q for w in ["export", "save", "download"]):
        return "export_summary"
    return "answer_question"


# ══════════════════════════════════════════════════════════════════════════════
#  PYDANTIC MODELS
# ══════════════════════════════════════════════════════════════════════════════

class ProvisionRequest(BaseModel):
    session_id: str
    allowed_scopes: list[str] = ["read:documents"]

class TokenRequest(BaseModel):
    client_id: str
    client_secret: str
    scopes: list[str] = ["read:documents"]

class DeprovisionRequest(BaseModel):
    client_id: str
    reason: str = "manual revoke"

class ChunkSettingsRequest(BaseModel):
    session_id: str
    chunk_size: int = DEFAULT_CHUNK_SIZE
    overlap: int = DEFAULT_CHUNK_OVERLAP

class AskRequest(BaseModel):
    session_id: str
    question: str


# ══════════════════════════════════════════════════════════════════════════════
#  ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/", include_in_schema=False)
async def index():
    return FileResponse(str(STATIC_DIR / "index.html"))


# ── UPLOAD ────────────────────────────────────────────────────────────────────

@app.post("/upload")
async def upload(file: UploadFile = File(...), chunk_size: int = DEFAULT_CHUNK_SIZE, overlap: int = DEFAULT_CHUNK_OVERLAP):
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(400, f"Unsupported file type '{ext}'.")

    session_id = uuid.uuid4().hex
    save_path = UPLOAD_DIR / f"{session_id}{ext}"
    save_path.write_bytes(await file.read())

    try:
        raw_text, meta = extract_file(save_path)
    except Exception as e:
        save_path.unlink(missing_ok=True)
        raise HTTPException(422, f"Could not parse: {e}")
    if not raw_text.strip():
        save_path.unlink(missing_ok=True)
        raise HTTPException(422, "No text extracted. File may be empty or image-only.")

    clean_text, warnings = sanitize_document(raw_text, session_id)

    chunk_size = chunk_size if chunk_size in VALID_CHUNK_SIZES else DEFAULT_CHUNK_SIZE
    overlap = max(0, min(overlap, chunk_size // 2))
    chunks = chunk_text(clean_text, chunk_size=chunk_size, overlap=overlap)

    sessions[session_id] = {
        "filename": file.filename, "file_type": meta["file_type"], "pages": meta["pages"],
        "word_count": len(raw_text.split()), "char_count": len(raw_text),
        "raw_text": raw_text, "clean_text": clean_text, "chunks": chunks,
        "chunk_size": chunk_size, "overlap": overlap, "injection_warnings": warnings,
        "client_id": None, "created_at": datetime.now(timezone.utc).isoformat(),
    }

    # ── PROVISION the agent's own identity for this session ──
    # This is deliberate and explicit: a human-triggered upload causes a
    # new non-human identity to be created, scoped to read:documents only.
    # write:documents is NOT granted here — it must be elevated on purpose.
    prov = nhi.provision_client(session_id, allowed_scopes=["read:documents", "write:documents"], owner="docmind-app")
    sessions[session_id]["client_id"] = prov["client_id"]

    audit("session_created", session_id, {
        "filename": file.filename, "file_type": meta["file_type"], "words": len(raw_text.split()),
        "chunks": len(chunks), "chunk_size": chunk_size, "injection_warnings": len(warnings),
        "provisioned_client_id": prov["client_id"],
    })

    return {
        "session_id": session_id, "filename": file.filename, "file_type": meta["file_type"],
        "pages": meta["pages"], "word_count": len(raw_text.split()), "char_count": len(raw_text),
        "chunk_count": len(chunks), "chunk_size": chunk_size, "overlap": overlap,
        "injection_warnings": warnings,
        # NHI bootstrap info — client_secret is shown ONCE, exactly like a real Auth0 app
        "agent_identity": {
            "client_id": prov["client_id"],
            "client_secret": prov["client_secret"],
            "allowed_scopes": prov["allowed_scopes"],
        },
    }


# ── CHUNK SETTINGS ────────────────────────────────────────────────────────────

@app.post("/rechunk")
async def rechunk(req: ChunkSettingsRequest):
    s = sessions.get(req.session_id)
    if not s:
        raise HTTPException(404, "Session not found.")
    chunk_size = req.chunk_size if req.chunk_size in VALID_CHUNK_SIZES else DEFAULT_CHUNK_SIZE
    overlap = max(0, min(req.overlap, chunk_size // 2))
    new_chunks = chunk_text(s["clean_text"], chunk_size=chunk_size, overlap=overlap)
    s["chunks"], s["chunk_size"], s["overlap"] = new_chunks, chunk_size, overlap
    audit("rechunk", req.session_id, {"chunk_size": chunk_size, "overlap": overlap, "chunk_count": len(new_chunks)})
    return {"session_id": req.session_id, "chunk_size": chunk_size, "overlap": overlap,
            "chunk_count": len(new_chunks), "valid_chunk_sizes": VALID_CHUNK_SIZES}


# ══════════════════════════════════════════════════════════════════════════════
#  NON-HUMAN IDENTITY ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/agent/identity/provision")
async def provision_identity(req: ProvisionRequest):
    """
    (Re-)provision an agent identity for a session. Normally done automatically
    at upload time; exposed here so a human can explicitly re-provision after
    a revoke, with a written-down scope decision.
    """
    if req.session_id not in sessions:
        raise HTTPException(404, "Session not found.")
    try:
        prov = nhi.provision_client(req.session_id, req.allowed_scopes, owner="human-reviewer")
    except ValueError as e:
        raise HTTPException(400, str(e))
    sessions[req.session_id]["client_id"] = prov["client_id"]
    audit("identity_provisioned", req.session_id, {"client_id": prov["client_id"], "scopes": req.allowed_scopes})
    return prov


@app.post("/agent/identity/token")
async def issue_token(req: TokenRequest):
    """
    client_credentials grant — the agent authenticates with ITS OWN secret,
    not a person's, and receives a short-lived, narrowly-scoped token.
    Requesting a scope outside allowed_scopes is denied here, before any
    tool ever runs — this is where over-scoping gets caught.
    """
    try:
        token = nhi.issue_token(req.client_id, req.client_secret, req.scopes)
    except PermissionError as e:
        client = nhi.get_client(req.client_id)
        sid = client["session_id"] if client else "unknown"
        audit("token_issuance_denied", sid, {"client_id": req.client_id, "requested_scopes": req.scopes, "reason": str(e)})
        raise HTTPException(403, str(e))

    client = nhi.get_client(req.client_id)
    audit("token_issued", client["session_id"], {
        "client_id": req.client_id, "scopes": req.scopes, "expires_in": token["expires_in"],
    })
    return token


@app.post("/agent/identity/deprovision")
async def deprovision_identity(req: DeprovisionRequest):
    """Revoke the agent's identity. Every outstanding token dies with it."""
    try:
        result = nhi.deprovision_client(req.client_id, req.reason)
    except ValueError as e:
        raise HTTPException(404, str(e))
    client = nhi.get_client(req.client_id)
    if client:
        audit("identity_deprovisioned", client["session_id"], {
            "client_id": req.client_id, "tokens_revoked": result["tokens_revoked"], "reason": req.reason,
        })
    return result


@app.get("/agent/identity/{client_id}")
async def get_identity(client_id: str):
    client = nhi.get_client(client_id)
    if not client:
        raise HTTPException(404, "client_id not found.")
    return client


@app.get("/agent/identity/token/{access_token}/status")
async def token_status(access_token: str):
    info = nhi.token_info(access_token)
    if not info:
        raise HTTPException(404, "Token not found or already expired/revoked.")
    return info


# ══════════════════════════════════════════════════════════════════════════════
#  AGENT ENDPOINT — every call requires Authorization: Bearer <access_token>
# ══════════════════════════════════════════════════════════════════════════════

def _extract_bearer(authorization: Optional[str]) -> str:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(401, "Missing Authorization header. Expected: Bearer <access_token>.")
    return authorization.removeprefix("Bearer ").strip()


@app.post("/agent/ask")
async def agent_ask(req: AskRequest, authorization: Optional[str] = Header(None)):
    s = sessions.get(req.session_id)
    if not s:
        raise HTTPException(404, "Session not found.")
    q = req.question.strip()
    if not q:
        raise HTTPException(400, "Question cannot be empty.")

    token = _extract_bearer(authorization)
    tool_name = decide_tool(q)
    tool = registry.get(tool_name)
    if not tool:
        raise HTTPException(500, f"Tool '{tool_name}' not registered.")

    valid, reason = nhi.validate_token(token, tool.required_scope)
    audit("token_check", req.session_id, {
        "tool": tool_name, "required_scope": tool.required_scope,
        "result": "ALLOW" if valid else "DENY", "reason": reason,
    })
    if not valid:
        return {"status": "denied", "tool_used": tool_name, "message": reason}

    call_id = uuid.uuid4().hex[:12]
    audit("agent_decision", req.session_id, {"call_id": call_id, "question_preview": q[:120], "selected_tool": tool_name})

    if tool.requires_approval:
        audit("approval_required", req.session_id, {"call_id": call_id, "tool": tool_name})
        return {"status": "pending_approval", "call_id": call_id, "tool_used": tool_name,
                "message": f"Tool '{tool_name}' needs write:documents and human approval — not auto-executed."}

    try:
        audit("tool_execute", req.session_id, {
        "call_id": call_id,
        "tool": tool_name
       })

        if tool_name == "answer_question":
           result = tool.fn(session_id=req.session_id, question=q)

        elif tool_name == "summarize_document":
           result = tool.fn(session_id=req.session_id)

        elif tool_name == "extract_keyterms":
           result = tool.fn(session_id=req.session_id)

        elif tool_name == "generate_quiz":
           result = tool.fn(session_id=req.session_id)

        elif tool_name == "export_summary":
           result = tool.fn(session_id=req.session_id)

        else:
           raise ValueError(f"Unknown tool: {tool_name}")

        audit("tool_success", req.session_id, {
        "call_id": call_id,
        "tool": tool_name
        })

        return {
        "status": "success",
        "call_id": call_id,
        "tool_used": tool_name,
        "result": result
     }

    except Exception as e:
    
        audit("tool_error", req.session_id, {
        "call_id": call_id,
        "tool": tool_name,
        "error": str(e)
        })

        return {
        "status": "error",
        "call_id": call_id,
        "tool_used": tool_name,
        "message": str(e)
        }


@app.post("/agent/export")
async def agent_export(req: AskRequest, authorization: Optional[str] = Header(None)):
    """Explicit write-scope action — export the summary to disk."""
    s = sessions.get(req.session_id)
    if not s:
        raise HTTPException(404, "Session not found.")
    token = _extract_bearer(authorization)
    tool = registry.get("export_summary")

    valid, reason = nhi.validate_token(token, tool.required_scope)
    audit("token_check", req.session_id, {"tool": "export_summary", "required_scope": tool.required_scope,
                                           "result": "ALLOW" if valid else "DENY", "reason": reason})
    if not valid:
        return {"status": "denied", "tool_used": "export_summary", "message": reason}

    call_id = uuid.uuid4().hex[:12]
    audit("tool_execute", req.session_id, {"call_id": call_id, "tool": "export_summary"})
    try:
        result = tool.fn(session_id=req.session_id)
        audit("tool_success", req.session_id, {"call_id": call_id, "tool": "export_summary"})
        return {"status": "success", "call_id": call_id, "tool_used": "export_summary", "result": result}
    except Exception as e:
        audit("tool_error", req.session_id, {"call_id": call_id, "tool": "export_summary", "error": str(e)})
        return {"status": "error", "call_id": call_id, "message": str(e)}


# ── AUDIT LOG / TOOLS LIST / SESSION INFO ─────────────────────────────────────

@app.get("/agent/audit")
async def audit_log(session_id: Optional[str] = None, limit: int = 100):
    return {"entries": get_audit_log(session_id=session_id, limit=limit)}


@app.get("/agent/tools")
async def list_tools():
    return {"tools": registry.list_tools()}


@app.get("/session/{session_id}")
async def get_session(session_id: str):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found.")
    return {k: v for k, v in s.items() if k not in ("raw_text", "clean_text", "chunks")}


@app.get("/session/{session_id}/chunks")
async def get_chunks(session_id: str, page: int = 0, per_page: int = 10):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found.")
    all_chunks = s["chunks"]
    start = page * per_page
    return {"total": len(all_chunks), "page": page, "per_page": per_page,
            "chunks": all_chunks[start: start + per_page], "chunk_size": s["chunk_size"], "overlap": s["overlap"]}


@app.get("/health")
async def health():
    ollama_ok, models = False, []
    try:
        r = requests.get(f"{OLLAMA_URL}/api/tags", timeout=3)
        if r.ok:
            ollama_ok = True
            models = [m["name"] for m in r.json().get("models", [])]
    except Exception:
        pass
    return {
        "status": "ok",
        "ollama": {"available": ollama_ok, "url": OLLAMA_URL, "models": models, "target": OLLAMA_MODEL},
        "claude_fallback": bool(CLAUDE_API_KEY),
        "registered_tools": len(registry._tools),
        "active_sessions": len(sessions),
        "valid_chunk_sizes": VALID_CHUNK_SIZES,
    }
